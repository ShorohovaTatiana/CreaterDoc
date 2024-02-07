from flask import Blueprint, request, jsonify, send_file
from docx import Document
from docx.shared import Cm
import openpyxl
import json
from io import BytesIO

fileDownload = Blueprint('fileDownload', __name__)

class SetEncoder(json.JSONEncoder):
    def default(self, obj):
        if isinstance(obj, set):
            return list(obj)
        return json.JSONEncoder.default(self, obj)


@fileDownload.route('/api/download', methods=['POST'])
def download_file():
    request_data = request.get_json()
    type = request_data['type'].strip()
    direction = request_data['direction'].strip()
    date = request_data['date']
    print(date)
    file = None
    if (type == 'enrollment' and direction == 'acquaintance'):
        file = openF(1, date)
    if (type == 'enrollment' and direction == 'analyze'):
        file = openF(2, date)
    if (type == 'deduction' and direction == 'acquaintance'):
        file = openF(3, date)
    if (type == 'deduction' and direction == 'analyze'):
        file = openF(4, date)
    if (type == 'officialMemo'):
        file = openF(5, date)
    if (file == None):
        return jsonify('По такому запросу не нашлось файлов'), 400

    f = BytesIO()
    document.save(f)
    length = f.tell()
    f.seek(0)


    return send_file(f, as_attachment=True, download_name='report.doc')
    # apiResponse = createApiResponse()
    # return apiResponse


# def createApiResponse():
#     bufferFile = writeBufferExcelFile()
#     mimetype = 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
#     return send_file(bufferFile, mimetype=mimetype)
#
#
# def writeBufferExcelFile():
#     buffer = BytesIO()
#     workbook = xlsxwriter.Workbook(buffer)
#     worksheet = workbook.add_worksheet()
#     worksheet.write(0, 0, "Hi There, its working!")
#     workbook.close()
#     buffer.seek(0)
#     return buffer

# - Это если захочешь эксель отправлять


a = []
document = Document()
wb = openpyxl.load_workbook(filename='files/318.xlsx')
worksheet = wb["Лист1"]

def list_students(c):
    temp = []
    buf1 = []
    buf2 = []
    for i in range(2, worksheet.max_row):
        temp.append(worksheet.cell(row=i, column=3).value)
        a.append(temp)
        if worksheet.cell(row=i, column=5).value == 'Знакомство с Python':
            buf1.append(worksheet.cell(row=i, column=3).value)
            a.append(buf1)
        if worksheet.cell(row=i, column=5).value == 'Python и анализ данных':
            buf2.append(worksheet.cell(row=i, column=3).value)
            a.append(buf2)

    temp = sorted(temp)
    buf1 = sorted(buf1)
    buf2 = sorted(buf2)
    chouse_docx(c, buf1, buf2, temp)



def chouse_docx(c, b1, b2, t):
    if c == 0:
        for i in range(0, len(b1)):
            p = document.add_paragraph(f' {b1[i].title()} ')
            p.style = 'List 3'
            pf = p.paragraph_format
            pf.left_indent = Cm(5.75)

    if c == 1:
        for i in range(0, len(b2)):
            p = document.add_paragraph(f' {b2[i].title()} ')
            p.style = 'List 3'
            pf = p.paragraph_format
            pf.left_indent = Cm(5.75)

    if c == 2:
        for i in range(0, len(t)):
            p = document.add_paragraph(f' {t[i].title()} ')
            p.style = 'List 3'
            pf = p.paragraph_format
            pf.left_indent = Cm(5.75)


def docx_creater(k, c, date):
    #date = list(d)
    #date[0],date[1],date[2],date[3],date[4],date[5],date[6],date[7],date[8],date[9] = date[9],date[8],date[7],date[6],date[5],date[4],date[3],date[2],date[1],date[0]
    #''.join(date)
    if c != 2:
        p = document.add_paragraph()
        p.paragraph_format.alignment = 1
        p.add_run('МИНОБРНАУКИ РОССИИ '
                  'ФЕДЕРАЛЬНОЕ ГОСУДАРСТВЕННОЕ БЮДЖЕТНОЕ ОБРАЗОВАТЕЛЬНОЕ УЧРЕЖДЕНИЕ '
                  'ВЫСШЕГО ОБРАЗОВАНИЯ '
                  '«ВОРОНЕЖСКИЙ ГОСУДАРСТВЕННЫЙ УНИВЕРСИТЕТ» '
                  '(ФГБОУ ВО «ВГУ»)').bold = True

        p = document.add_paragraph()
        p.paragraph_format.alignment = 1
        p.add_run('Приказ').bold = True

        p = document.add_paragraph(f'{date}	Воронеж   	№ 3-2820')
        p.paragraph_format.alignment = 1

        p = document.add_paragraph()
        p.add_run(f'    По личному составу обучающихся по программе дополнительного образования')
        if k == 0:
            in_out = 'Зачисление'
        if k == 1:
            in_out = 'Отчисление'
        p.add_run(f'       ({in_out})')

        document.add_paragraph(
            f'    В соответствии с Порядком организации и осуществления деятельности при сетевой форме реализации '
            f'образовательных программ (утвержден приказами Министерства науки и высшего образования и  Министерства '
            f'просвещения Российской Федерации от 5 августа 2020 года  №882/391)  и договором оферты о сетевой форме '
            f'реализации дополнительных общеразвивающих программ между Автономной некоммерческой организацией дополнительного'
            f' школьного образования «Школа Анализа данных» и федеральным государственным бюджетным образовательным '
            f'учреждением высшего образования «Воронежский государственный университет» (ВГУ) от 23 августа 2022 года. ')
        document.add_paragraph(f'   п р и к а з ы в а ю:')


        if c == 0:
            name_pr = 'Знакомство с Python'
        if c == 1:
            name_pr = 'Python и анализ данных'
        document.add_paragraph(
            f'    1. {in_out} с {date} в число обучающихся факультета компьютерных наук очной формы обучения по '
            f'дополнительной общеразвивающей программе «{name_pr}»:')

        list_students(c)

        document.add_paragraph(f"Первый проректор – проректор по учебной работе               "
                               f"                                  ФИО")
        document.add_paragraph(f'ПРОЕКТ ВНОСИТ – ')
        document.add_paragraph(f'Декан факультета компьютерных наук                      '
                               f'                                             ФИО __.__.20__')
        document.add_paragraph(f'СОГЛАСОВАНО:')
        document.add_paragraph(f'Начальник ПФО                                                               '
                               f'                                                    ФИО __.__.20__')
        document.add_paragraph(f'Ведущий экономист ПФО                                            '
                               f'                                                    ФИО __.__.20__')
        document.add_paragraph(f'Директор ИДПО                                                           '
                               f'                                                        ФИО __.__.20__')
        document.add_paragraph(f'')
        document.add_paragraph(f'Расчет рассылки: ИДПО, ПФО, бухгалтерия (3), ф-т компьютерных наук')
    else:
        document.add_paragraph(f'Факультет компьютерных наук	Ректору ВГУ Ендовицкому Д.А.')
        document.add_paragraph(f'')
        document.add_paragraph(f'СЛУЖЕБНАЯ ЗАПИСКА')
        document.add_paragraph(f'от {date}  № ______')
        document.add_paragraph(f'    В соответствии с планом довузовской работы ФКН и соглашением, заключенным между '
                               f'ФГБОУ ВО «ВГУ» и компанией Яндекс, на факультете компьютерных наук с {date}'
                               f' по 30.06.2022 будут проводиться очные занятия для школьников по программе Лицея '
                               f'Академии Яндекса второго года обучения. Прошу разрешить пропуск учащихся в '
                               f'соответствии со списком с {date}:')
        list_students(k)
        document.add_paragraph(f'    Обязуюсь осуществлять контроль за соблюдением санитарно-эпидемиологических '
                               f'правил указанными учащимися в условиях предупреждения распространения новой '
                               f'коронавирусной инфекции.')
        document.add_paragraph(f'')
        document.add_paragraph(f'')
        document.add_paragraph(f'Декан ФКН								                  ФИО')

    return document



def openF(n, date):
    if (n == 1):
        return docx_creater(0, 0, date)  # Зачисление _ Знакомство с Python
    if (n == 2):
        return docx_creater(0, 1, date)  # Зачисление _ Python и анализ данных
    if (n == 3):
        return docx_creater(1, 0, date)  # Отчисление _ Знакомство с Python
    if (n == 4):
        return docx_creater(1, 1, date)  # Отчисление _ Python и анализ данных
    if (n == 5):
        return docx_creater(2, 2, date)  # Записка
